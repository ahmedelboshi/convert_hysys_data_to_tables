import openpyxl
from docx import Document

# the project aim to convert the excel file to tables in word file

# Open the XLSX file
file_name= 'H2'
workbook = openpyxl.load_workbook(f'data/{file_name}.xlsx')
# first select material streams sheet
mete_sheet = workbook['Material Streams']
# read the first row start by third column and add values to a list
materials = [cell.value for cell in mete_sheet[1]][2::1]
print(materials)
# read the first column start by second row and add values to a list
streams = [cell.value for cell in mete_sheet['A']][2::1]
print(streams)

# go through each column and add it is values to list then add the list to a list
data=[]
for i in range(2,mete_sheet.max_column+1):
    data.append([cell.value for cell in mete_sheet[i]][2::1])
print(data)


# Iterate through the streams
# for row in range(2, sheet.max_row+1):
#     stream_name = sheet.cell(row=row, column=1).value
    
#     # Create a new doc for the stream
#     document = Document()
#     document.add_heading(stream_name, 0)
    
#     # Add the properties to the doc
#     properties = sheet.cell(row=1, column=2).value
#     document.add_paragraph(properties)
    
#     # Add a table to the doc
#     table = document.add_table(rows=1, cols=2)
#     table.style = 'Table Grid'
    
#     # Add the property and unit to the table
#     for col in range(2, sheet.max_column+1):
#         property_cell = sheet.cell(row=1, column=col)
#         unit_cell = sheet.cell(row=row, column=col)
#         table.cell(0, 0).text = property_cell.value
#         table.cell(0, 1).text = unit_cell.value
    
    # Save the doc
# document.save(f'res/{file_name}.docx')
