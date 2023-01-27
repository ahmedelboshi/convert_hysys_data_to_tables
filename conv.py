import pandas as pd
import docx
# Read the sheet named "Sheet1" from the excel file
file_name= 'R'
df = pd.read_excel(f'data/{file_name}.xlsx', sheet_name='Material Streams')

materials = list(df.columns[2::1])

# Create an empty list to store the columns
columns_list = []
# Iterate through each column in the dataframe
for column in df.columns:
    # Add the column data to the list
    columns_list.append(list(df[column]))
properties = columns_list[0][1:]

unites_list = columns_list[1][1:]
# replace none values with empty string
unites_list[0] = ''

# go through each column and add it is values and remove the first  values
columns_list = [column[1::1] for column in columns_list[2::1]]

# go through each column and add unit to it from unites_list
data=[]
for i,col in enumerate(columns_list):
    data.append([f"{round(col[i],3)} {unites_list[i]}" for i in range(len(col))])

# write  list dictionary containing properties and data

results={}
for index,material_data in enumerate(data):
    # compaine properties and data in list of dictionaries {properties:data}
    
    results[materials[index]] = {"properties":[f"{properties[i]}: {material_data[i]}" for i in range(len(material_data))]}
    # for i,z in enumerate(d):


# now we have a dictionary containing materials and properties and data
# we need to get materiels compassion from Compositions sheet
# Read the sheet named "Compositions" from the excel file

df = pd.read_excel(f'data/{file_name}.xlsx', sheet_name='Compositions')

# Get the list of components
components = [list(df[column])for column in df.columns][0][1:]
# remove "Comp Mole Frac (" and ")" from each component
components = [component.replace("Comp Mole Frac (", "").replace(")", "") for component in components] 

# remove the first two items in data
data=[ list(df[column])[1:] for column in df.columns][2::1]

# go through each item in data and add it to a dictionary with component as key
# and value as the data
for i,stream in enumerate(results):
    d = {}

    for index,c in enumerate(components):
        d[c]=data[i][index]
    results[stream]["components"] = d


# now we have a dictionary containing materials and properties and data
# we need create 2 tables for each stream first table for properties and second for components

document = docx.Document()
start=0
for stream in results:
    # Create a new doc for the stream
    # document.add_heading(material, 0)
    
    # Add the properties to the doc
    properties = results[stream]["properties"]
    # make paragraph centered
    paragraph = document.add_paragraph()
    paragraph.alignment = 1
    paragraph.add_run(f"{stream} Properties  Table 3.{start}.1").bold = True
    # 
    # document.add_paragraph("Properties")
    
    # Add a table to the doc
    table = document.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
   
    # Add the property and unit to the table
    for i in range(len(properties)):
        row = table.add_row()
        table.cell(i, 0).text = properties[i].split(":")[0]
        # add padding to cell top


        table.cell(i, 1).text = properties[i].split(":")[1]

    # add space between tables
    paragraph = document.add_paragraph()
    paragraph.alignment = 1
    paragraph.add_run(" ").bold = True
    paragraph = document.add_paragraph()
    paragraph.alignment = 1
    paragraph.add_run(" ").bold = True
    


    paragraph = document.add_paragraph()
    paragraph.alignment = 1
    
    paragraph.add_run(f"{stream} Components Table 3.{start}.1  ").bold = True

    # Add a table to the doc
    table = document.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    
    # Add the property and unit to the table
    # Add the components to the doc
    components = results[stream]["components"]
    total=1
    for i,component in enumerate(components):
        row = table.add_row()
        table.cell(i, 0).text = component
        # if item is last in list 
        if i == len(components)-1:
            # if total is less than zero make if zero
            if total < 0:
                total=0
            table.cell(i, 1).text = str(round(total,4))
        else:
            total-=round(float(components[component]),4)
            table.cell(i, 1).text = str(round(float(components[component]),4))
    
    # add a page break
    document.add_page_break()
    # Save the doc
    start+=1
document.save(f'res/{file_name}.docx')

print("done")